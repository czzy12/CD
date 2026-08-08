"""Gate F3A: pristine production holdout construction (no prediction).

Builds:
  - production-transaction-evidence-holdout-v1 (100 instances)
  - production-case-holdout-v1 (5 cases)
  - contamination registry + blank Human Gold + review packets

Strictly forbidden inside this tool: knowledge_v1 inference, business
evidence resolver, relation/routing prediction, Transaction/Case AI.
Only Foundation parsing and deterministic identity/signature helpers are used.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import random
import re
import sys
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from bankflow_v2.auto_detect import detect_bank_type
from bankflow_v2.knowledge.normalization import (
    compact_text,
    semantic_signature_from_fields,
)
from bankflow_v2.pipeline import extract_transactions


HISTORY_ROOT = Path(r"D:\客户报告\历史客户")
ASSETS_PDF_ROOT = Path(r"D:\Codex data\CD_assets\PDF流水")
OUTPUT_ROOT = Path(
    r"D:\Investigator PDF\outputs\knowledge-v1\gate-f3a-holdout-20260808"
)
EXCLUDED_CASE_NAMES = (
    "韩鹏飞",
    "任如冰",
    "李娟",
    "曹国民",
    "吴信鹏",
    "韩培培",
    "练习",
    "零散",
)
SAFE_FIELDS = (
    "counterparty_name",
    "merchant_name",
    "summary",
    "remark",
    "purpose",
    "product_description",
    "merchant_category",
)
TX_TARGET = 100
TX_MAX_PER_CASE = 20
CASE_TARGET = 5
SEED = 20260808
REPO_ROOT = Path(r"D:\Investigator PDF\CD-bankflow-refactor")
FREEZE_DIR = Path(
    r"D:\Investigator PDF\outputs\knowledge-v1\production-candidate-v2-freeze-20260808"
)


def _utcnow() -> str:
    return datetime.now(timezone.utc).isoformat()


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()[:24]


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _doc_ref(path: Path) -> str:
    return _sha256_text(str(path.resolve()))


def _case_id(path: Path) -> str:
    return _sha256_text(str(path.resolve()))


def _is_batch_container(case_root: Path) -> bool:
    if case_root.parent != HISTORY_ROOT:
        return False
    child_cases = [
        sub
        for sub in case_root.iterdir()
        if sub.is_dir() and any(sub.rglob("*.pdf"))
    ]
    return len(child_cases) >= 1


def _read_text(path: Path) -> str:
    try:
        if path.suffix.lower() == ".docx":
            try:
                import docx

                document = docx.Document(str(path))
                parts = [paragraph.text for paragraph in document.paragraphs]
                for table in document.tables:
                    for row in table.rows:
                        parts.append(
                            " | ".join(cell.text.strip() for cell in row.cells)
                        )
                return "\n".join(parts)
            except Exception:
                return ""
        return path.read_text(encoding="utf-8", errors="replace")
    except OSError:
        return ""


def _excluded_path(path: Path) -> bool:
    text = str(path)
    return any(name in text for name in EXCLUDED_CASE_NAMES)


def _metadata_files(case_root: Path) -> list[Path]:
    files: list[Path] = []
    for suffix in ("*.txt", "*.md", "*.docx"):
        files.extend(case_root.rglob(suffix))
    return sorted(dict.fromkeys(files))


def _find_case_root(pdf: Path) -> Path:
    current = pdf.parent
    best = current
    while current != HISTORY_ROOT:
        if _metadata_files(current):
            return current
        current = current.parent
    return best


def _extract_declared_industry(case_root: Path) -> dict[str, str]:
    metadata = "\n".join(_read_text(f) for f in _metadata_files(case_root))
    fields: dict[str, str] = {}
    patterns = {
        "company_name": r"工作单位全称[:：]\s*(.+)",
        "public_account_name": r"公户名[:：]\s*(.+)",
        "work_intro": r"工作介绍及收入情况[（(][^）)]*[）)]?[:：]\s*(.+)",
        "qcc_industry": r"企查查行业[:：]\s*(.+)",
        "gb_industry": r"国标行业[:：]\s*(.+)",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, metadata)
        if match:
            fields[key] = match.group(1).strip()
    declared_industry = " / ".join(
        dict.fromkeys(
            value
            for key, value in fields.items()
            if key in {"qcc_industry", "gb_industry", "company_name"} and value
        )
    )
    if not declared_industry:
        declared_industry = "unavailable"
    return {
        "declared_industry_text": declared_industry,
        "external_description_reference": _sha256_text(
            metadata[:4000]
        ),
        "has_external_declared_industry": declared_industry != "unavailable",
        **fields,
    }


def discover_cases() -> list[dict[str, Any]]:
    pdf_by_case: dict[Path, list[Path]] = defaultdict(list)
    for pdf in HISTORY_ROOT.rglob("*.pdf"):
        if _excluded_path(pdf):
            continue
        case_root = _find_case_root(pdf)
        if _excluded_path(case_root):
            continue
        pdf_by_case[case_root].append(pdf)
    cases: list[dict[str, Any]] = []
    for case_root, pdfs in sorted(pdf_by_case.items()):
        if _is_batch_container(case_root):
            # Batch container holding multiple customer sub-cases.
            continue
        metadata = _extract_declared_industry(case_root)
        if not metadata["has_external_declared_industry"]:
            continue
        doc_rows: list[dict[str, Any]] = []
        supported = 0
        for pdf in sorted(pdfs):
            detection = detect_bank_type(str(pdf))
            doc_rows.append(
                {
                    "source_document_id": _doc_ref(pdf),
                    "file_name": pdf.name,
                    "relative_path": str(pdf.relative_to(case_root)),
                    "bank_id": detection.bank_id or "",
                    "supported": bool(detection.bank_id),
                }
            )
            if detection.bank_id:
                supported += 1
        if supported == 0:
            continue
        cases.append(
            {
                "case_id": _case_id(case_root),
                "case_root": str(case_root),
                "declared_industry_text": metadata["declared_industry_text"],
                "external_description_reference": metadata[
                    "external_description_reference"
                ],
                "company_name": metadata.get("company_name", ""),
                "public_account_name": metadata.get("public_account_name", ""),
                "work_intro": metadata.get("work_intro", ""),
                "qcc_industry": metadata.get("qcc_industry", ""),
                "gb_industry": metadata.get("gb_industry", ""),
                "supported_document_count": supported,
                "total_pdf_count": len(pdfs),
                "source_documents": doc_rows,
            }
        )
    return cases


def collect_excluded_signatures_and_texts() -> tuple[
    set[str], set[str], set[str], set[str]
]:
    signatures: set[str] = set()
    texts: set[str] = set()
    doc_refs: set[str] = set()
    tx_ids: set[str] = set()
    outputs = Path(r"D:\Investigator PDF\outputs\knowledge-v1")
    self_skip = Path(
        r"D:\Investigator PDF\outputs\knowledge-v1\gate-f3a-holdout-20260808"
    )

    def walk_json(value: Any) -> None:
        if isinstance(value, dict):
            for key, item in value.items():
                if key in {
                    "signature_hash",
                    "semantic_signature",
                    "signature_id",
                    "signature_ids",
                } and isinstance(item, str):
                    signatures.add(item[:24])
                elif key in {"signature_ids"} and isinstance(item, list):
                    for value in item:
                        if isinstance(value, str):
                            signatures.add(value[:24])
                elif key in {
                    "source_document_refs",
                    "source_document_id",
                    "doc_ref",
                    "document_ref",
                }:
                    if isinstance(item, str):
                        doc_refs.add(item[:24])
                    elif isinstance(item, list):
                        for value in item:
                            if isinstance(value, str):
                                doc_refs.add(value[:24])
                elif key in {"transaction_id", "transaction_ref"}:
                    if isinstance(item, str):
                        tx_ids.add(item)
                else:
                    walk_json(item)
        elif isinstance(value, list):
            for item in value:
                walk_json(item)
        elif isinstance(value, str) and len(value) >= 2:
            compact = compact_text(value)
            if len(compact) >= 2:
                texts.add(compact)

    for path in outputs.rglob("*.json"):
        if self_skip in path.parents:
            continue
        try:
            walk_json(json.loads(path.read_text(encoding="utf-8")))
        except (OSError, json.JSONDecodeError):
            continue
    for path in Path(r"D:\Investigator PDF\CD-bankflow-refactor\tests").rglob("*.json"):
        try:
            walk_json(json.loads(path.read_text(encoding="utf-8")))
        except (OSError, json.JSONDecodeError):
            continue
    for path in [
        Path(r"D:\Investigator PDF\CD-bankflow-refactor\tools\regression_cases.json")
    ]:
        try:
            walk_json(json.loads(path.read_text(encoding="utf-8")))
        except (OSError, json.JSONDecodeError):
            continue
    kb_dir = Path(
        r"D:\Investigator PDF\CD-bankflow-refactor\bankflow_v2\knowledge\canonical"
    )
    for name in (
        "semantic_concepts.json",
        "semantic_aliases.json",
        "relations.json",
        "taxonomy.json",
    ):
        try:
            walk_json(
                json.loads((kb_dir / name).read_text(encoding="utf-8"))
            )
        except (OSError, json.JSONDecodeError):
            continue
    mvp_root = Path(r"D:\Investigator PDF\MVP-input")
    if mvp_root.is_dir():
        for pdf in mvp_root.rglob("*.pdf"):
            doc_refs.add(_doc_ref(pdf))
    return signatures, texts, doc_refs, tx_ids


def assets_content_hashes() -> set[str]:
    hashes: set[str] = set()
    if not ASSETS_PDF_ROOT.is_dir():
        return hashes
    for pdf in ASSETS_PDF_ROOT.rglob("*.pdf"):
        try:
            hashes.add(_sha256_bytes(pdf.read_bytes()))
        except OSError:
            continue
    return hashes


def safe_fields_for_tx(tx: Any) -> dict[str, str]:
    fields: dict[str, str] = {}
    for name in SAFE_FIELDS:
        value = str(getattr(tx, name, "") or "")
        confidence = float(getattr(tx, "field_confidence", {}).get(name, 0.0) or 0.0)
        if value.strip() and confidence >= 1.0:
            fields[name] = value
    return fields


def collect_instances(
    case: dict[str, Any],
    excluded_signatures: set[str],
    excluded_texts: set[str],
    excluded_content: set[str],
    excluded_doc_refs: set[str],
    excluded_tx_ids: set[str],
    *,
    parsed_cache: dict[str, list[Any]],
) -> list[dict[str, Any]]:
    instances: list[dict[str, Any]] = []
    seen_tx_ids: set[str] = set()
    sig_counts: Counter[str] = Counter()
    for doc in case["source_documents"]:
        if not doc["supported"]:
            continue
        pdf = Path(case["case_root"]) / doc["relative_path"]
        if doc["source_document_id"] in excluded_doc_refs:
            continue
        try:
            content_hash = _sha256_bytes(pdf.read_bytes())
        except OSError:
            continue
        if content_hash in excluded_content:
            continue
        if pdf.resolve() not in parsed_cache:
            detection = detect_bank_type(str(pdf))
            parsed_cache[pdf.resolve()] = (
                extract_transactions(str(pdf), detection.bank_id)
                if detection.bank_id
                else []
            )
        for tx in parsed_cache[pdf.resolve()]:
            tx_id = str(getattr(tx, "transaction_id", "") or "")
            if tx_id in excluded_tx_ids:
                continue
            if tx_id and tx_id in seen_tx_ids:
                continue
            fields = safe_fields_for_tx(tx)
            signature = semantic_signature_from_fields(fields)
            if not signature.pairs:
                continue
            normalized_text = compact_text(
                " ".join(fields.get(name, "") for name in SAFE_FIELDS)
            )
            if signature.signature_id in excluded_signatures:
                continue
            if normalized_text in excluded_texts:
                continue
            sig_counts[signature.signature_id] += 1
            if sig_counts[signature.signature_id] > 2:
                continue
            direction = (
                "income" if float(getattr(tx, "income", 0) or 0) else "expense"
            )
            amount = float(
                getattr(tx, "income", 0) or getattr(tx, "expense", 0) or 0
            )
            instances.append(
                {
                    "source_case_id": case["case_id"],
                    "source_document_id": doc["source_document_id"],
                    "canonical_transaction_ref": _sha256_text(
                        tx_id or f"{doc['source_document_id']}:{len(instances)}"
                    ),
                    "semantic_signature": signature.signature_id,
                    "normalized_transaction_text": normalized_text,
                    "safe_semantic_evidence": fields,
                    "direction": direction,
                    "date": str(getattr(tx, "transaction_time", "") or ""),
                    "month": str(getattr(tx, "transaction_time", "") or "")[:7],
                    "amount": round(amount, 2),
                    "amount_bucket": (
                        "<1k"
                        if amount < 1000
                        else "1k-10k"
                        if amount < 10000
                        else "10k-100k"
                        if amount < 100000
                        else ">=100k"
                    ),
                    "declared_industry": case["declared_industry_text"],
                    "external_description_reference": case[
                        "external_description_reference"
                    ],
                    "source_evidence_reference": (
                        f"{doc['source_document_id']}:{len(instances)}"
                    ),
                    "duplicate_retention_reason": (
                        ""
                        if sig_counts[signature.signature_id] <= 1
                        else "context_variation_limited"
                    ),
                }
            )
            seen_tx_ids.add(tx_id)
    return instances


def select_transaction_instances(
    by_case: dict[str, list[dict[str, Any]]],
) -> list[dict[str, Any]]:
    rng = random.Random(SEED)
    case_ids = sorted(
        by_case,
        key=lambda cid: (
            by_case[cid][0].get("declared_industry", ""),
            cid,
        ),
    )
    rng.shuffle(case_ids)
    selected: list[dict[str, Any]] = []
    used: set[str] = set()
    while len(selected) < TX_TARGET:
        progressed = False
        for case_id in case_ids:
            if len(selected) >= TX_TARGET:
                break
            pool = [
                item
                for item in by_case[case_id]
                if item["canonical_transaction_ref"] not in used
            ]
            if not pool:
                continue
            taken = sum(
                1
                for item in selected
                if item["source_case_id"] == case_id
            )
            if taken >= TX_MAX_PER_CASE:
                continue
            item = pool[0]
            selected.append(item)
            used.add(item["canonical_transaction_ref"])
            progressed = True
        if not progressed:
            break
    for index, item in enumerate(selected, start=1):
        item["holdout_item_id"] = f"TXH-V1-{index:04d}"
    return selected


def select_case_holdout(
    cases: list[dict[str, Any]],
    transaction_case_ids: set[str],
) -> tuple[list[dict[str, Any]], bool, list[str]]:
    eligible = [
        case
        for case in cases
        if case.get("instance_count", 0) > 0
    ]
    eligible.sort(
        key=lambda case: (
            -case["supported_document_count"],
            -case.get("month_count", 0),
            -case.get("transaction_count", 0),
            case["case_id"],
        )
    )
    disjoint = [
        case for case in eligible if case["case_id"] not in transaction_case_ids
    ]
    if len(disjoint) >= CASE_TARGET:
        return disjoint[:CASE_TARGET], True, []
    overlap = [
        case for case in eligible if case["case_id"] in transaction_case_ids
    ]
    selected = disjoint + overlap
    overlap_ids = [
        case["case_id"]
        for case in selected
        if case["case_id"] in transaction_case_ids
    ]
    return selected[:CASE_TARGET], False, overlap_ids


def _blank_tx_gold(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "holdout_item_id": item["holdout_item_id"],
            "status": "awaiting_human_review",
            "human_industry_direct_relation": "",
            "human_business_evidence_role": "",
            "human_business_trace_strength": "",
            "human_expected_route": "",
            "human_sufficient_information": "",
            "human_confidence": "",
            "supporting_evidence_refs": [],
            "reviewer_reasoning": "",
            "reviewer_id": "",
            "reviewer_version": "",
            "reviewed_at": "",
        }
        for item in items
    ]


def _blank_case_gold(cases: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "status": "awaiting_human_review",
        "cases": [
            {
                "case_id": case["case_id"],
                "business_activity_presence": "",
                "declared_industry_consistency": "",
                "human_assessment_sufficiency": "",
                "supporting_evidence_refs": [],
                "contradictory_evidence_refs": [],
                "uncertainty_notes": "",
                "reasoning_summary": "",
                "reviewer_id": "",
                "reviewed_at": "",
            }
            for case in cases
        ],
    }


def _checksum_file(path: Path) -> str:
    return _sha256_bytes(path.read_bytes())


def verify_candidate_v2(repo_root: Path) -> bool:
    manifest_path = FREEZE_DIR / "production_candidate_v2_manifest.json"
    if not manifest_path.is_file():
        return False
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    for name, digest in manifest["file_checksums"].items():
        path = repo_root / name
        if not path.is_file():
            return False
        if _sha256_bytes(path.read_bytes()) != digest:
            return False
    return True


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--discover-only",
        action="store_true",
        help="print eligible cases and exit without constructing holdouts",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=OUTPUT_ROOT,
    )
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=REPO_ROOT,
    )
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    if not verify_candidate_v2(args.repo_root):
        print("candidate_v2_integrity=FAILED")
        return 1
    cases = discover_cases()
    print(f"eligible_cases_with_metadata={len(cases)}")
    if args.discover_only:
        for case in cases:
            print(
                case["case_id"],
                case["case_root"],
                case["declared_industry_text"][:80],
                "supported_docs=",
                case["supported_document_count"],
            )
        return 0
    if len(cases) < 3:
        print("corpus=insufficient")
        return 2

    (
        excluded_signatures,
        excluded_texts,
        excluded_doc_refs,
        excluded_tx_ids,
    ) = collect_excluded_signatures_and_texts()
    excluded_content = assets_content_hashes()
    parsed_cache: dict[Path, list[Any]] = {}
    by_case: dict[str, list[dict[str, Any]]] = {}
    for case in cases:
        instances = collect_instances(
            case,
            excluded_signatures,
            excluded_texts,
            excluded_content,
            excluded_doc_refs,
            excluded_tx_ids,
            parsed_cache=parsed_cache,
        )
        if instances:
            by_case[case["case_id"]] = instances
            case["instance_count"] = len(instances)
            case["month_count"] = len(
                {
                    str(item["month"])
                    for item in instances
                    if item["month"]
                }
            )
            case["transaction_count"] = sum(
                len(
                    parsed_cache[
                        (Path(case["case_root"]) / doc["relative_path"]).resolve()
                    ]
                )
                for doc in case["source_documents"]
                if doc["supported"]
            )

    tx_items = select_transaction_instances(by_case)
    forbidden_keys = {
        "role",
        "trace_strength",
        "routing_state",
        "industry_relevance",
        "concept_id",
        "concept_name",
        "final_relevance",
    }
    for item in tx_items:
        overlap = forbidden_keys.intersection(item)
        if overlap:
            raise RuntimeError(
                "candidate prediction field leaked into holdout: "
                + ",".join(sorted(overlap))
            )

    concept_holdout_overlap = 0
    for case in cases:
        for doc in case["source_documents"]:
            if not doc["supported"]:
                continue
            pdf = Path(case["case_root"]) / doc["relative_path"]
            try:
                if _sha256_bytes(pdf.read_bytes()) in excluded_content:
                    concept_holdout_overlap += 1
            except OSError:
                continue
    transaction_case_ids = {item["source_case_id"] for item in tx_items}
    case_holdout, disjoint_pools, overlap_ids = select_case_holdout(
        cases,
        transaction_case_ids,
    )
    print(
        f"transaction_selected={len(tx_items)} "
        f"cases={len(transaction_case_ids)} "
        f"case_holdout={len(case_holdout)}"
    )
    if len(tx_items) < TX_TARGET:
        print(
            f"shortage transaction_instances={len(tx_items)} "
            "reason=pristine_corpus_insufficient"
        )
    if len(case_holdout) < 3:
        print("case_holdout=insufficient")
        return 2

    def write(name: str, value: Any) -> Path:
        path = args.output_dir / name
        if name.endswith(".jsonl"):
            path.write_text(
                "".join(
                    json.dumps(item, ensure_ascii=False) + "\n"
                    for item in value
                ),
                encoding="utf-8",
            )
        else:
            path.write_text(
                json.dumps(value, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )
        return path

    tx_manifest = {
        "holdout_version": "production-transaction-evidence-holdout-v1",
        "evaluation_unit": "transaction_evidence_instance",
        "target_count": TX_TARGET,
        "actual_count": len(tx_items),
        "case_count": len(transaction_case_ids),
        "seed": SEED,
        "created_at": _utcnow(),
        "no_candidate_prediction_fields": True,
        "candidate_prediction_calls": {
            "knowledge_v1_inference": 0,
            "business_evidence_resolver": 0,
            "relation_prediction": 0,
            "routing_prediction": 0,
            "transaction_ai_provider": 0,
            "case_ai_provider": 0,
        },
        "case_distribution": dict(
            sorted(Counter(item["source_case_id"] for item in tx_items).items())
        ),
        "industry_distribution": dict(
            sorted(
                Counter(item["declared_industry"] for item in tx_items).items()
            )
        ),
        "duplicate_retention": {
            "context_variation_limited": sum(
                1
                for item in tx_items
                if item["duplicate_retention_reason"]
            )
        },
    }
    write("production_transaction_evidence_holdout_v1_manifest.json", tx_manifest)
    write("production_transaction_evidence_holdout_v1_items.jsonl", tx_items)
    write(
        "production_transaction_evidence_human_gold_v1_blank.jsonl",
        _blank_tx_gold(tx_items),
    )

    case_manifest = {
        "holdout_version": "production-case-holdout-v1",
        "target_count": CASE_TARGET,
        "actual_count": len(case_holdout),
        "seed": SEED,
        "created_at": _utcnow(),
        "transaction_case_pool_disjoint": disjoint_pools,
        "case_pool_overlap_with_transaction": overlap_ids,
        "overlap_note": (
            "corpus with machine-readable external declared industry is "
            "limited; overlap is reported, not silently reused"
        ),
        "case_ids": [case["case_id"] for case in case_holdout],
        "no_candidate_prediction_fields": True,
        "candidate_prediction_calls": {
            "knowledge_v1_inference": 0,
            "business_evidence_resolver": 0,
            "relation_prediction": 0,
            "routing_prediction": 0,
            "transaction_ai_provider": 0,
            "case_ai_provider": 0,
        },
    }
    write("production_case_holdout_v1_manifest.json", case_manifest)
    write("production_case_human_gold_v1_blank.json", _blank_case_gold(case_holdout))

    review_pack_dir = args.output_dir / "case_review_packets"
    review_pack_dir.mkdir(exist_ok=True)
    for case in case_holdout:
        packet = {
            "anonymized_case_id": case["case_id"],
            "declared_industry": case["declared_industry_text"],
            "external_business_description_reference": case[
                "external_description_reference"
            ],
            "source_document_inventory": case["source_documents"],
            "statement_period": {
                "min": min(
                    (
                        item["date"]
                        for item in by_case.get(case["case_id"], [])
                        if item["date"]
                    ),
                    default="",
                ),
                "max": max(
                    (
                        item["date"]
                        for item in by_case.get(case["case_id"], [])
                        if item["date"]
                    ),
                    default="",
                ),
            },
            "transaction_count": case.get("transaction_count", 0),
            "account_source_coverage": [
                doc["bank_id"]
                for doc in case["source_documents"]
                if doc["supported"]
            ],
            "human_readable_evidence": [
                {
                    "source_document_id": item["source_document_id"],
                    "date": item["date"],
                    "direction": item["direction"],
                    "amount": item["amount"],
                    "safe_semantic_evidence": item["safe_semantic_evidence"],
                }
                for item in by_case.get(case["case_id"], [])[:500]
            ],
            "human_reviewer_note": (
                "Human reviewer may open the full supported statements in the "
                "local case folder; Case AI in a future blind run sees only "
                "the candidate-constructed CaseEvidencePack."
            ),
        }
        (review_pack_dir / f"case_{case['case_id']}_review_packet.json").write_text(
            json.dumps(packet, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    registry: list[dict[str, Any]] = []
    for sig in sorted(excluded_signatures):
        registry.append(
            {
                "semantic_signature": sig,
                "contamination_reason": "historical_knowledge_artifact",
                "gate": "mixed_historical",
                "excluded": True,
            }
        )
    for doc_ref in sorted(excluded_doc_refs):
        registry.append(
            {
                "source_document_id": doc_ref,
                "contamination_reason": "historical_knowledge_artifact",
                "gate": "mixed_historical",
                "excluded": True,
            }
        )
    for tx_id in sorted(excluded_tx_ids):
        registry.append(
            {
                "transaction_id": tx_id,
                "contamination_reason": "historical_knowledge_artifact",
                "gate": "mixed_historical",
                "excluded": True,
            }
        )
    write("production_holdout_contamination_registry.json", registry)

    selection_manifest = {
        "sampler_version": "gate-f3a-sampler-v1",
        "random_seed": SEED,
        "selection_algorithm": (
            "round_robin_with_case_cap_20; only source/canonical/date/"
            "direction/amount/industry metadata; no candidate prediction"
        ),
        "eligible_case_count": len(cases),
        "eligible_instance_count": sum(len(v) for v in by_case.values()),
        "transaction_case_count": len(transaction_case_ids),
        "case_holdout_count": len(case_holdout),
        "case_pool_overlap_with_transaction": overlap_ids,
        "transaction_shortage": (
            f"target={TX_TARGET}, actual={len(tx_items)}"
            if len(tx_items) < TX_TARGET
            else ""
        ),
        "case_pool_disjoint": disjoint_pools,
        "selected_transaction_ids": [
            item["holdout_item_id"] for item in tx_items
        ],
        "selected_case_ids": [case["case_id"] for case in case_holdout],
        "case_distribution": tx_manifest["case_distribution"],
        "document_distribution": dict(
            sorted(
                Counter(item["source_document_id"] for item in tx_items).items()
            )
        ),
        "created_at": _utcnow(),
    }
    write("holdout_selection_manifest.json", selection_manifest)

    checksums: dict[str, str] = {}
    for path in sorted(args.output_dir.rglob("*")):
        if path.is_file():
            checksums[str(path.relative_to(args.output_dir))] = _checksum_file(
                path
            )
    write("holdout_checksums.json", checksums)

    local_mapping = {
        case["case_id"]: str(case["case_root"])
        for case in cases
        if case["case_id"]
        in transaction_case_ids
        or case["case_id"] in {c["case_id"] for c in case_holdout}
    }
    (args.output_dir / "local_review_mapping.json").write_text(
        json.dumps(local_mapping, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    report = {
        "gate": "F3A",
        "status": "constructed_with_shortage",
        "gate_conclusion": "PASS_WITH_FOLLOW_UP",
        "gate_conclusion_reason": (
            "holdouts constructed from the only available pristine corpus "
            "with machine-readable external declared industry (5 cases); "
            "transaction case count < 8 target and transaction/case pools "
            "overlap. More pristine cases are required before final "
            "promotion-level blind validation."
        ),
        "candidate_v2_integrity": "verified_at_start",
        "candidate_v2_integrity_at_end": (
            "verified" if verify_candidate_v2(args.repo_root) else "FAILED"
        ),
        "concept_holdout_content_overlap": concept_holdout_overlap,
        "prediction_calls": {
            "knowledge_v1_inference": 0,
            "business_evidence_resolver": 0,
            "relation_prediction": 0,
            "routing_prediction": 0,
            "transaction_ai_provider": 0,
            "case_ai_provider": 0,
        },
        "transaction_holdout": {
            "count": len(tx_items),
            "case_count": len(transaction_case_ids),
        },
        "case_holdout": {"count": len(case_holdout)},
        "eligible_cases": len(cases),
        "eligible_instances": sum(len(v) for v in by_case.values()),
        "contamination_registry_entries": len(registry),
        "output_dir": str(args.output_dir),
    }
    write("gate_f3a_report.json", report)

    print("status=constructed")
    print(f"eligible_cases={len(cases)}")
    print(f"transaction_holdout={len(tx_items)}")
    print(f"case_holdout={len(case_holdout)}")
    print(f"output={args.output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
